#!/usr/bin/env python3
# pip install python-docx lxml
from pathlib import Path
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
import re, sys, unicodedata, os

###############################################################################
# Utility helpers
###############################################################################

def slugify(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")

def clean_text(text: str) -> str:
    text = unicodedata.normalize("NFKC", text or "")
    text = text.replace("\u00A0", " ")
    # Keep accented letters, strip control chars
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", "", text)
    return text.strip()

def is_list_paragraph(paragraph):
    p_style = (paragraph.style.name or "").lower()
    return "bullet" in p_style or "number" in p_style

def is_bullet(paragraph):
    return "bullet" in (paragraph.style.name or "").lower()

def is_numbered(paragraph):
    return "number" in (paragraph.style.name or "").lower()

def build_paragraph_html(paragraph, rels):
    """Rebuilds paragraph preserving links, images, and text order."""
    for element in paragraph._element.iter():
        # Hyperlink detection
        if element.tag == qn("w:hyperlink"):
            rid = element.get(qn("r:id"))
            url = None
            if rid and rid in rels and rels[rid].reltype == RT.HYPERLINK:
                url = rels[rid].target_ref

            link_text_parts = []
            for t in element.findall(".//w:t", namespaces=element.nsmap):
                link_text_parts.append(clean_text(t.text))
            link_text = "".join(link_text_parts)

            if url and link_text:
                yield {"text": f'<a href="{url}" target="_blank" rel="noopener">{link_text}</a>'}

            # Prevent child <w:t> from being processed again
            for t in element.findall(".//w:t", namespaces=element.nsmap):
                t.text = None
            continue

        # Image detection
        if element.tag == qn("a:blip"):
            rid = element.get(qn("r:embed"))
            if rid and rid in rels and rels[rid].reltype == RT.IMAGE:
                yield {"image_rid": rid}
            continue

        # Plain text
        if element.tag == qn("w:t") and element.text:
            yield {"text": clean_text(element.text)}

###############################################################################
# Paths / wrapper templates
###############################################################################

ARTICLE_DIR = Path("articles")
TOP_WRAPPER = Path("article-wrapper-top.html").read_text(encoding="utf-8")
BOT_WRAPPER = Path("article-wrapper-bottom.html").read_text(encoding="utf-8")

###############################################################################
# Main conversion loop
###############################################################################

for docx_file in ARTICLE_DIR.glob("*.docx"):
    if docx_file.name.startswith(("~$", ".")):
        continue

    try:
        doc         = Document(docx_file)
        doc_title   = clean_text(doc.core_properties.title or (doc.paragraphs[0].text if doc.paragraphs else docx_file.stem))
        slug        = slugify(doc_title) or docx_file.stem
        html_path   = ARTICLE_DIR / f"{slug}.html"
        image_dir   = Path("assets/images") / slug
        image_dir.mkdir(parents=True, exist_ok=True)

        html_parts  = []
        image_count = 0
        rels        = doc.part._rels

        html_parts.append(f"<h1>{doc_title}</h1>\n")

        list_mode = None

        for para in doc.paragraphs:
            if not para.text and not para.runs:
                continue

            if is_list_paragraph(para):
                if is_bullet(para):
                    if list_mode != "ul":
                        if list_mode:
                            html_parts.append(f"</{list_mode}>\n")
                        html_parts.append("<ul>\n")
                        list_mode = "ul"
                elif is_numbered(para):
                    if list_mode != "ol":
                        if list_mode:
                            html_parts.append(f"</{list_mode}>\n")
                        html_parts.append("<ol>\n")
                        list_mode = "ol"

                li_content = []
                for chunk in build_paragraph_html(para, rels):
                    if "text" in chunk:
                        li_content.append(chunk["text"])
                    elif "image_rid" in chunk:
                        image_count += 1
                        image_part = rels[chunk["image_rid"]].target_part
                        ext = os.path.splitext(image_part.partname)[1] or ".png"
                        img_name = f"{slug}-img{image_count}{ext}"
                        (image_dir / img_name).write_bytes(image_part.blob)
                        li_content.append(
                            f'<img src="../assets/images/{slug}/{img_name}" '
                            f'alt="Image {image_count}" style="max-width:600px;display:block;margin:1rem auto;border-radius:8px;" />'
                        )
                html_parts.append(f"<li>{''.join(li_content)}</li>\n")

            else:
                if list_mode:
                    html_parts.append(f"</{list_mode}>\n")
                    list_mode = None

                p_content = []
                for chunk in build_paragraph_html(para, rels):
                    if "text" in chunk:
                        p_content.append(chunk["text"])
                    elif "image_rid" in chunk:
                        image_count += 1
                        image_part = rels[chunk["image_rid"]].target_part
                        ext = os.path.splitext(image_part.partname)[1] or ".png"
                        img_name = f"{slug}-img{image_count}{ext}"
                        (image_dir / img_name).write_bytes(image_part.blob)
                        p_content.append(
                            f'<img src="../assets/images/{slug}/{img_name}" '
                            f'alt="Image {image_count}" style="max-width:600px;display:block;margin:1rem auto;border-radius:8px;" />'
                        )
                if p_content:
                    html_parts.append(f"<p>{''.join(p_content)}</p>\n")

        if list_mode:
            html_parts.append(f"</{list_mode}>\n")

        full_html = f"{TOP_WRAPPER}\n<article>\n{''.join(html_parts)}</article>\n{BOT_WRAPPER}"
        html_path.write_text(full_html, encoding="utf-8")
        print(f"✅ Converted: {docx_file.name} → {html_path.name} ({image_count} image{'s' if image_count != 1 else ''})")

    except Exception as e:
        print(f"❌ Error processing {docx_file.name}: {e}", file=sys.stderr)
